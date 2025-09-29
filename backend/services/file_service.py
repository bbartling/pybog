"""
File management service with hybrid storage and EventBus integration.
Handles file uploads, storage decisions (BYTEA vs file_path), and state management.
"""

import asyncio
import logging
import os
import uuid
import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Any, List, Optional, BinaryIO

from fastapi import UploadFile
import PyPDF2
from docx import Document
import magic

from core.database import DatabaseManager, get_database
from core.events import EventBus, Event
from core.config import get_file_retention_config
from models.file_models import (
    FileRecord, FileType, ProgressState, StorageType, 
    FileUploadRequest, FileStateUpdate, FileMetadata, 
    FileListResponse, FileCleanupResult, TextExtractionResult, FilePreview
)

logger = logging.getLogger(__name__)


class FileService:
    """
    File management service with EventBus integration for event emission.
    
    Handles:
    - File uploads with automatic BYTEA (<10MB) vs file_path (>=10MB) decision
    - File retrieval for both storage types
    - File state management with ProgressState transitions
    - Event emission for all file operations
    """
    
    def __init__(self, event_bus: EventBus, db_manager: Optional[DatabaseManager] = None):
        self.event_bus = event_bus
        self.db_manager = db_manager
        self.config = get_file_retention_config()
        self.file_size_limit_bytes = self.config.max_file_size_mb * 1024 * 1024
        
        # Create upload directory for large files
        self.upload_dir = Path("data/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"FileService initialized with {self.config.max_file_size_mb}MB BYTEA limit")
    
    async def _get_db(self) -> DatabaseManager:
        """Get database manager instance."""
        if self.db_manager:
            return self.db_manager
        return await get_database()
    
    async def _emit_event(self, session_id: str, operation: str, event_type: str, data: Dict[str, Any]) -> None:
        """Emit an event to the EventBus."""
        try:
            event = Event(
                type=event_type,
                session_id=session_id,
                operation=operation,
                data=data
            )
            await self.event_bus.publish(session_id, event)
            logger.debug(f"Emitted {event_type} event for session {session_id}, operation {operation}")
        except Exception as e:
            logger.error(f"Failed to emit event: {e}")
    
    async def upload_file(self, session_id: str, file: UploadFile, file_type: FileType = FileType.UPLOAD) -> FileRecord:
        """
        Upload a file with automatic BYTEA (<10MB) vs file_path (>=10MB) decision.
        
        Args:
            session_id: Session identifier
            file: FastAPI UploadFile object
            file_type: Type of file being uploaded
            
        Returns:
            FileRecord with storage information
            
        Raises:
            ValueError: If file is invalid or too large
            Exception: If upload fails
        """
        try:
            # Emit upload started event
            await self._emit_event(
                session_id, 
                "upload", 
                "progress", 
                {
                    "state": ProgressState.QUEUED,
                    "message": f"Starting upload of {file.filename}",
                    "filename": file.filename
                }
            )
            
            # Read file content
            file_content = await file.read()
            file_size = len(file_content)
            
            # Validate file size (max 50MB total limit)
            max_total_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_total_size:
                error_msg = f"File too large: {file_size} bytes (max {max_total_size} bytes)"
                await self._emit_event(
                    session_id,
                    "upload",
                    "error",
                    {
                        "error_code": "FILE",
                        "message": error_msg,
                        "filename": file.filename
                    }
                )
                raise ValueError(error_msg)
            
            # Generate unique filename
            file_extension = Path(file.filename).suffix if file.filename else ""
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            
            # Determine storage method
            use_bytea = file_size < self.file_size_limit_bytes
            storage_type = StorageType.BYTEA if use_bytea else StorageType.FILE_PATH
            
            await self._emit_event(
                session_id,
                "upload", 
                "progress",
                {
                    "state": ProgressState.PROCESSING,
                    "message": f"Storing file ({storage_type.value}): {file.filename}",
                    "filename": file.filename,
                    "storage_type": storage_type.value,
                    "file_size": file_size
                }
            )
            
            db = await self._get_db()
            
            if use_bytea:
                # Store in database as BYTEA
                file_record = await self._store_bytea_file(
                    db, session_id, unique_filename, file.filename, 
                    file.content_type, file_type, file_content, file_size
                )
            else:
                # Store on filesystem with path reference
                file_path = await self._store_file_on_disk(file_content, unique_filename)
                file_record = await self._store_file_path_reference(
                    db, session_id, unique_filename, file.filename,
                    file.content_type, file_type, file_path, file_size
                )
            
            # Update state to complete
            await self.update_file_state(file_record.id, ProgressState.COMPLETE)
            file_record.state = ProgressState.COMPLETE
            file_record.storage_type = storage_type
            
            await self._emit_event(
                session_id,
                "upload",
                "progress", 
                {
                    "state": ProgressState.COMPLETE,
                    "message": f"Upload completed: {file.filename}",
                    "filename": file.filename,
                    "file_id": file_record.id,
                    "storage_type": storage_type.value,
                    "file_size": file_size
                }
            )
            
            logger.info(f"File uploaded successfully: {file.filename} -> {unique_filename} ({storage_type.value})")
            return file_record
            
        except Exception as e:
            logger.error(f"File upload failed for {file.filename}: {e}")
            
            await self._emit_event(
                session_id,
                "upload",
                "error",
                {
                    "error_code": "FILE",
                    "message": f"Upload failed: {str(e)}",
                    "filename": file.filename if file else "unknown"
                }
            )
            raise
    
    async def _store_bytea_file(
        self, db: DatabaseManager, session_id: str, filename: str, 
        original_name: str, mime_type: Optional[str], file_type: FileType,
        file_content: bytes, file_size: int
    ) -> FileRecord:
        """Store file content as BYTEA in database."""
        
        query = """
            INSERT INTO files (session_id, filename, original_name, mime_type, file_type, 
                             file_data, file_size, state, created_at)
            VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9)
            RETURNING id, session_id, filename, original_name, mime_type, file_type, 
                     file_size, state, metadata, created_at, archived_at
        """
        
        now = datetime.now(timezone.utc)
        row = await db.fetch_one(
            query, session_id, filename, original_name, mime_type, 
            file_type.value, file_content, file_size, ProgressState.PROCESSING.value, now
        )
        
        # Convert JSON string back to dict for Pydantic model
        import json
        file_data = dict(row)
        if isinstance(file_data.get('metadata'), str):
            file_data['metadata'] = json.loads(file_data['metadata'])
        elif file_data.get('metadata') is None:
            file_data['metadata'] = {}
        return FileRecord(**file_data)
    
    async def _store_file_on_disk(self, file_content: bytes, filename: str) -> str:
        """Store file content on disk and return the path."""
        file_path = self.upload_dir / filename
        
        # Ensure directory exists
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Write file content
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return str(file_path)
    
    async def _store_file_path_reference(
        self, db: DatabaseManager, session_id: str, filename: str,
        original_name: str, mime_type: Optional[str], file_type: FileType,
        file_path: str, file_size: int
    ) -> FileRecord:
        """Store file path reference in database."""
        
        query = """
            INSERT INTO files (session_id, filename, original_name, mime_type, file_type,
                             file_path, file_size, state, created_at)
            VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9)
            RETURNING id, session_id, filename, original_name, mime_type, file_type,
                     file_size, state, metadata, created_at, archived_at
        """
        
        now = datetime.now(timezone.utc)
        row = await db.fetch_one(
            query, session_id, filename, original_name, mime_type,
            file_type.value, file_path, file_size, ProgressState.PROCESSING.value, now
        )
        
        # Convert JSON string back to dict for Pydantic model
        import json
        file_data = dict(row)
        if isinstance(file_data.get('metadata'), str):
            file_data['metadata'] = json.loads(file_data['metadata'])
        elif file_data.get('metadata') is None:
            file_data['metadata'] = {}
        return FileRecord(**file_data)    

    async def get_file_data(self, file_id: int) -> bytes:
        """
        Retrieve file data for both BYTEA and file_path storage types.
        
        Args:
            file_id: File identifier
            
        Returns:
            File content as bytes
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file has no data
        """
        try:
            db = await self._get_db()
            
            # Get file metadata and determine storage type
            query = """
                SELECT id, session_id, filename, file_data, file_path, file_size, state
                FROM files WHERE id = $1
            """
            row = await db.fetch_one(query, file_id)
            
            if not row:
                raise FileNotFoundError(f"File with ID {file_id} not found")
            
            # Check if file has BYTEA data
            if row['file_data'] is not None:
                logger.debug(f"Retrieved BYTEA file data for file {file_id}")
                return bytes(row['file_data'])
            
            # Check if file has file_path
            elif row['file_path'] is not None:
                file_path = Path(row['file_path'])
                if not file_path.exists():
                    raise FileNotFoundError(f"File not found on disk: {file_path}")
                
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                logger.debug(f"Retrieved file_path file data for file {file_id}")
                return content
            
            else:
                raise ValueError(f"File {file_id} has no data (neither BYTEA nor file_path)")
                
        except Exception as e:
            logger.error(f"Failed to retrieve file data for {file_id}: {e}")
            raise
    
    async def get_file_metadata(self, file_id: int) -> FileMetadata:
        """
        Get file metadata including computed storage type.
        
        Args:
            file_id: File identifier
            
        Returns:
            FileMetadata object
            
        Raises:
            FileNotFoundError: If file doesn't exist
        """
        try:
            db = await self._get_db()
            
            query = """
                SELECT id, session_id, filename, original_name, mime_type, file_type,
                       file_data, file_path, file_size, state, metadata, created_at, archived_at,
                       get_file_storage_type(id) as storage_type
                FROM files WHERE id = $1
            """
            row = await db.fetch_one(query, file_id)
            
            if not row:
                raise FileNotFoundError(f"File with ID {file_id} not found")
            
            # Convert row to FileMetadata with JSON handling
            import json
            file_metadata_dict = dict(row)
            if isinstance(file_metadata_dict.get('metadata'), str):
                file_metadata_dict['metadata'] = json.loads(file_metadata_dict['metadata'])
            elif file_metadata_dict.get('metadata') is None:
                file_metadata_dict['metadata'] = {}
            
            metadata = FileMetadata(
                id=file_metadata_dict['id'],
                session_id=file_metadata_dict['session_id'],
                filename=file_metadata_dict['filename'],
                original_name=file_metadata_dict['original_name'],
                mime_type=file_metadata_dict['mime_type'],
                file_type=FileType(file_metadata_dict['file_type']),
                file_size=file_metadata_dict['file_size'],
                state=ProgressState(file_metadata_dict['state']),
                storage_type=StorageType(file_metadata_dict['storage_type']),
                metadata=file_metadata_dict['metadata'],
                created_at=file_metadata_dict['created_at'],
                archived_at=file_metadata_dict['archived_at']
            )
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to get file metadata for {file_id}: {e}")
            raise
    
    async def update_file_state(self, file_id: int, state: ProgressState, 
                               metadata: Optional[Dict[str, Any]] = None,
                               error_message: Optional[str] = None) -> bool:
        """
        Update file state with ProgressState enum transitions.
        
        Args:
            file_id: File identifier
            state: New progress state
            metadata: Optional metadata to update
            error_message: Optional error message for failed state
            
        Returns:
            True if update was successful
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If state transition is invalid
        """
        try:
            db = await self._get_db()
            
            # Get current file info
            current_file = await self.get_file_metadata(file_id)
            
            # Validate state transition
            self._validate_state_transition(current_file.state, state)
            
            # Build update query
            update_parts = ["state = $2"]
            params = [file_id, state.value]
            param_count = 2
            
            if metadata is not None:
                param_count += 1
                update_parts.append(f"metadata = ${param_count}")
                import json
                params.append(json.dumps(metadata))
            
            # Add error message to metadata if provided
            if error_message and state == ProgressState.FAILED:
                if metadata is None:
                    metadata = current_file.metadata.copy()
                metadata['error_message'] = error_message
                
                if "metadata = $" not in " ".join(update_parts):
                    param_count += 1
                    update_parts.append(f"metadata = ${param_count}")
                    import json
                    params.append(json.dumps(metadata))
            
            query = f"""
                UPDATE files 
                SET {', '.join(update_parts)}
                WHERE id = $1
            """
            
            result = await db.execute_query(query, *params)
            
            # Emit state change event
            await self._emit_event(
                current_file.session_id,
                "file_state_update",
                "progress",
                {
                    "state": state.value,
                    "message": f"File state updated to {state.value}",
                    "file_id": file_id,
                    "filename": current_file.filename,
                    "previous_state": current_file.state.value,
                    "error_message": error_message
                }
            )
            
            logger.info(f"File {file_id} state updated: {current_file.state.value} -> {state.value}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to update file state for {file_id}: {e}")
            raise
    
    def _validate_state_transition(self, current_state: ProgressState, new_state: ProgressState) -> None:
        """
        Validate that a state transition is allowed.
        
        Valid transitions:
        - QUEUED -> PROCESSING, FAILED
        - PROCESSING -> FINALIZING, COMPLETE, FAILED  
        - FINALIZING -> COMPLETE, FAILED
        - COMPLETE -> (no transitions allowed)
        - FAILED -> QUEUED (retry)
        """
        valid_transitions = {
            ProgressState.QUEUED: [ProgressState.PROCESSING, ProgressState.FAILED],
            ProgressState.PROCESSING: [ProgressState.FINALIZING, ProgressState.COMPLETE, ProgressState.FAILED],
            ProgressState.FINALIZING: [ProgressState.COMPLETE, ProgressState.FAILED],
            ProgressState.COMPLETE: [ProgressState.PROCESSING, ProgressState.FAILED],  # Allow reprocessing
            ProgressState.FAILED: [ProgressState.QUEUED, ProgressState.PROCESSING]  # Allow retry
        }
        
        if new_state not in valid_transitions.get(current_state, []):
            raise ValueError(
                f"Invalid state transition: {current_state.value} -> {new_state.value}. "
                f"Valid transitions from {current_state.value}: {[s.value for s in valid_transitions.get(current_state, [])]}"
            )
    
    async def list_session_files(self, session_id: str, file_type: Optional[FileType] = None) -> FileListResponse:
        """
        List all files for a session with optional filtering by file type.
        
        Args:
            session_id: Session identifier
            file_type: Optional file type filter
            
        Returns:
            FileListResponse with file metadata list
        """
        try:
            db = await self._get_db()
            
            # Build query with optional file type filter
            where_clause = "WHERE session_id = $1"
            params = [session_id]
            
            if file_type:
                where_clause += " AND file_type = $2"
                params.append(file_type.value)
            
            query = f"""
                SELECT id, session_id, filename, original_name, mime_type, file_type,
                       file_size, state, metadata, created_at, archived_at,
                       get_file_storage_type(id) as storage_type
                FROM files 
                {where_clause}
                ORDER BY created_at DESC
            """
            
            rows = await db.fetch_all(query, *params)
            
            # Convert rows to FileMetadata objects with JSON handling
            import json
            files = []
            for row in rows:
                file_metadata_dict = dict(row)
                if isinstance(file_metadata_dict.get('metadata'), str):
                    file_metadata_dict['metadata'] = json.loads(file_metadata_dict['metadata'])
                elif file_metadata_dict.get('metadata') is None:
                    file_metadata_dict['metadata'] = {}
                
                file_metadata = FileMetadata(
                    id=file_metadata_dict['id'],
                    session_id=file_metadata_dict['session_id'],
                    filename=file_metadata_dict['filename'],
                    original_name=file_metadata_dict['original_name'],
                    mime_type=file_metadata_dict['mime_type'],
                    file_type=FileType(file_metadata_dict['file_type']),
                    file_size=file_metadata_dict['file_size'],
                    state=ProgressState(file_metadata_dict['state']),
                    storage_type=StorageType(file_metadata_dict['storage_type']),
                    metadata=file_metadata_dict['metadata'],
                    created_at=file_metadata_dict['created_at'],
                    archived_at=file_metadata_dict['archived_at']
                )
                files.append(file_metadata)
            
            return FileListResponse(
                files=files,
                total_count=len(files),
                session_id=session_id
            )
            
        except Exception as e:
            logger.error(f"Failed to list files for session {session_id}: {e}")
            raise
    
    async def cleanup_old_files(self) -> FileCleanupResult:
        """
        Run file cleanup operations (archive and purge old files).
        
        Returns:
            FileCleanupResult with counts of processed files
        """
        try:
            db = await self._get_db()
            
            # Use database functions for cleanup
            archived_count = await db.fetch_val(
                "SELECT archive_old_bytea_files($1)", 
                self.config.archive_threshold_days
            )
            
            purged_count = await db.fetch_val(
                "SELECT purge_archived_files($1)",
                self.config.purge_threshold_days
            )
            
            result = FileCleanupResult(
                archived_count=archived_count or 0,
                purged_count=purged_count or 0
            )
            
            logger.info(f"File cleanup completed: {result.archived_count} archived, {result.purged_count} purged")
            
            # Emit cleanup completion event (no specific session)
            await self._emit_event(
                "system",
                "file_cleanup",
                "progress",
                {
                    "state": ProgressState.COMPLETE.value,
                    "message": f"File cleanup completed: {result.archived_count} archived, {result.purged_count} purged",
                    "archived_count": result.archived_count,
                    "purged_count": result.purged_count
                }
            )
            
            return result
            
        except Exception as e:
            logger.error(f"File cleanup failed: {e}")
            
            await self._emit_event(
                "system",
                "file_cleanup", 
                "error",
                {
                    "error_code": "DB",
                    "message": f"File cleanup failed: {str(e)}"
                }
            )
            raise
    
    async def delete_file(self, file_id: int) -> bool:
        """
        Delete a file and its associated data.
        
        Args:
            file_id: File identifier
            
        Returns:
            True if deletion was successful
            
        Raises:
            FileNotFoundError: If file doesn't exist
        """
        try:
            # Get file metadata first
            file_metadata = await self.get_file_metadata(file_id)
            
            db = await self._get_db()
            
            # If file is stored on disk, delete the physical file
            if file_metadata.storage_type == StorageType.FILE_PATH:
                file_path_query = "SELECT file_path FROM files WHERE id = $1"
                file_path_row = await db.fetch_one(file_path_query, file_id)
                
                if file_path_row and file_path_row['file_path']:
                    file_path = Path(file_path_row['file_path'])
                    if file_path.exists():
                        file_path.unlink()
                        logger.debug(f"Deleted physical file: {file_path}")
            
            # Delete database record
            delete_query = "DELETE FROM files WHERE id = $1"
            await db.execute_query(delete_query, file_id)
            
            # Emit deletion event
            await self._emit_event(
                file_metadata.session_id,
                "file_delete",
                "progress",
                {
                    "state": ProgressState.COMPLETE.value,
                    "message": f"File deleted: {file_metadata.filename}",
                    "file_id": file_id,
                    "filename": file_metadata.filename
                }
            )
            
            logger.info(f"File {file_id} ({file_metadata.filename}) deleted successfully")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete file {file_id}: {e}")
            raise
    
    async def extract_text_content(self, file_id: int) -> TextExtractionResult:
        """
        Extract text content from PDF, DOCX, or text files.
        
        Args:
            file_id: File identifier
            
        Returns:
            TextExtractionResult with extracted text and metadata
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported for text extraction
        """
        try:
            # Get file metadata and content
            file_metadata = await self.get_file_metadata(file_id)
            file_content = await self.get_file_data(file_id)
            
            # Emit extraction started event
            await self._emit_event(
                file_metadata.session_id,
                "text_extraction",
                "progress",
                {
                    "state": ProgressState.PROCESSING,
                    "message": f"Starting text extraction from {file_metadata.filename}",
                    "file_id": file_id,
                    "filename": file_metadata.filename
                }
            )
            
            # Determine extraction method based on file type and content
            extraction_method = self._determine_extraction_method(file_metadata, file_content)
            
            # Extract text based on method
            if extraction_method == "pdf":
                result = await self._extract_pdf_text(file_id, file_content, file_metadata)
            elif extraction_method == "docx":
                result = await self._extract_docx_text(file_id, file_content, file_metadata)
            elif extraction_method == "text":
                result = await self._extract_text_file(file_id, file_content, file_metadata)
            else:
                # Fallback to text extraction attempt
                result = await self._extract_fallback_text(file_id, file_content, file_metadata)
            
            # Update file metadata with extraction results
            extraction_metadata = {
                "text_extracted": True,
                "extraction_method": result.extraction_method,
                "word_count": result.word_count,
                "character_count": result.character_count,
                "page_count": result.page_count,
                "extracted_text": result.extracted_text  # Store extracted text in metadata
            }
            
            # Update metadata with extraction results (don't change state if already complete)
            if file_metadata.state != ProgressState.COMPLETE:
                await self.update_file_state(
                    file_id, 
                    ProgressState.COMPLETE, 
                    {**file_metadata.metadata, **extraction_metadata}
                )
            else:
                # Just update metadata without changing state - use direct database update
                db = await self._get_db()
                import json
                await db.execute_query(
                    "UPDATE files SET metadata = $1 WHERE id = $2",
                    json.dumps({**file_metadata.metadata, **extraction_metadata}),
                    file_id
                )
            
            # Emit extraction completed event
            await self._emit_event(
                file_metadata.session_id,
                "text_extraction",
                "progress",
                {
                    "state": ProgressState.COMPLETE,
                    "message": f"Text extraction completed: {result.word_count} words extracted",
                    "file_id": file_id,
                    "filename": file_metadata.filename,
                    "word_count": result.word_count,
                    "extraction_method": result.extraction_method
                }
            )
            
            logger.info(f"Text extraction completed for file {file_id}: {result.word_count} words")
            return result
            
        except Exception as e:
            logger.error(f"Text extraction failed for file {file_id}: {e}")
            
            # Update file state to failed
            await self.update_file_state(
                file_id, 
                ProgressState.FAILED, 
                error_message=f"Text extraction failed: {str(e)}"
            )
            
            # Emit extraction failed event
            await self._emit_event(
                file_metadata.session_id,
                "text_extraction",
                "error",
                {
                    "error_code": "FILE",
                    "message": f"Text extraction failed: {str(e)}",
                    "file_id": file_id,
                    "filename": file_metadata.filename
                }
            )
            
            # Return failed result instead of raising
            return TextExtractionResult(
                file_id=file_id,
                extracted_text="",
                word_count=0,
                character_count=0,
                extraction_method="failed",
                success=False,
                error_message=str(e)
            )
    
    def _determine_extraction_method(self, file_metadata: FileMetadata, file_content: bytes) -> str:
        """
        Determine the best extraction method based on file metadata and content.
        
        Args:
            file_metadata: File metadata
            file_content: File content bytes
            
        Returns:
            Extraction method: "pdf", "docx", "text", or "fallback"
        """
        # Check MIME type first
        if file_metadata.mime_type:
            if "pdf" in file_metadata.mime_type.lower():
                return "pdf"
            elif "word" in file_metadata.mime_type.lower() or "docx" in file_metadata.mime_type.lower():
                return "docx"
            elif "text" in file_metadata.mime_type.lower():
                return "text"
        
        # Check file extension
        file_ext = Path(file_metadata.filename).suffix.lower()
        if file_ext == ".pdf":
            return "pdf"
        elif file_ext in [".docx", ".doc"]:
            return "docx"
        elif file_ext in [".txt", ".md", ".csv", ".log"]:
            return "text"
        
        # Use python-magic to detect file type from content
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
            if "pdf" in mime_type.lower():
                return "pdf"
            elif "word" in mime_type.lower() or "docx" in mime_type.lower():
                return "docx"
            elif "text" in mime_type.lower():
                return "text"
        except Exception as e:
            logger.warning(f"Failed to detect MIME type with python-magic: {e}")
        
        # Default to fallback method
        return "fallback"
    
    async def _extract_pdf_text(self, file_id: int, file_content: bytes, file_metadata: FileMetadata) -> TextExtractionResult:
        """Extract text from PDF file using PyPDF2."""
        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            extracted_text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    extracted_text += page_text + "\n\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # Clean up extracted text
            extracted_text = extracted_text.strip()
            word_count = len(extracted_text.split()) if extracted_text else 0
            character_count = len(extracted_text)
            
            return TextExtractionResult(
                file_id=file_id,
                extracted_text=extracted_text,
                page_count=page_count,
                word_count=word_count,
                character_count=character_count,
                extraction_method="pdf",
                metadata={
                    "pdf_pages": page_count,
                    "pdf_encrypted": pdf_reader.is_encrypted
                }
            )
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_id: int, file_content: bytes, file_metadata: FileMetadata) -> TextExtractionResult:
        """Extract text from DOCX file using python-docx."""
        try:
            docx_stream = io.BytesIO(file_content)
            doc = Document(docx_stream)
            
            extracted_text = ""
            paragraph_count = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        extracted_text += " | ".join(row_text) + "\n"
            
            # Clean up extracted text
            extracted_text = extracted_text.strip()
            word_count = len(extracted_text.split()) if extracted_text else 0
            character_count = len(extracted_text)
            
            return TextExtractionResult(
                file_id=file_id,
                extracted_text=extracted_text,
                word_count=word_count,
                character_count=character_count,
                extraction_method="docx",
                metadata={
                    "docx_paragraphs": paragraph_count,
                    "docx_tables": table_count
                }
            )
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_text_file(self, file_id: int, file_content: bytes, file_metadata: FileMetadata) -> TextExtractionResult:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            extracted_text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    extracted_text = file_content.decode(encoding)
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if extracted_text is None:
                # Fallback to utf-8 with error handling
                extracted_text = file_content.decode('utf-8', errors='replace')
                encoding_used = 'utf-8-replace'
            
            # Clean up text
            extracted_text = extracted_text.strip()
            word_count = len(extracted_text.split()) if extracted_text else 0
            character_count = len(extracted_text)
            
            # Count lines
            line_count = len(extracted_text.splitlines()) if extracted_text else 0
            
            return TextExtractionResult(
                file_id=file_id,
                extracted_text=extracted_text,
                word_count=word_count,
                character_count=character_count,
                extraction_method="text",
                metadata={
                    "encoding_used": encoding_used,
                    "line_count": line_count
                }
            )
            
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise ValueError(f"Failed to extract text from text file: {str(e)}")
    
    async def _extract_fallback_text(self, file_id: int, file_content: bytes, file_metadata: FileMetadata) -> TextExtractionResult:
        """Fallback text extraction for unknown file types."""
        try:
            # Try to decode as text with error handling
            extracted_text = file_content.decode('utf-8', errors='replace')
            
            # Filter out non-printable characters except whitespace
            import string
            printable_chars = string.printable
            filtered_text = ''.join(char for char in extracted_text if char in printable_chars)
            
            # Clean up text
            filtered_text = filtered_text.strip()
            word_count = len(filtered_text.split()) if filtered_text else 0
            character_count = len(filtered_text)
            
            return TextExtractionResult(
                file_id=file_id,
                extracted_text=filtered_text,
                word_count=word_count,
                character_count=character_count,
                extraction_method="fallback",
                metadata={
                    "original_size": len(file_content),
                    "filtered_size": len(filtered_text),
                    "warning": "Fallback extraction used - results may be incomplete"
                }
            )
            
        except Exception as e:
            logger.error(f"Fallback text extraction failed: {e}")
            raise ValueError(f"Failed to extract text using fallback method: {str(e)}")
    
    async def generate_file_preview(self, file_id: int, max_preview_chars: int = 500) -> FilePreview:
        """
        Generate a file preview with extracted text content.
        
        Args:
            file_id: File identifier
            max_preview_chars: Maximum characters for preview text
            
        Returns:
            FilePreview with preview text and metadata
            
        Raises:
            FileNotFoundError: If file doesn't exist
        """
        try:
            # Get file metadata
            file_metadata = await self.get_file_metadata(file_id)
            
            # Check if text has already been extracted
            if file_metadata.metadata.get("text_extracted"):
                # Use cached extraction results
                full_text = file_metadata.metadata.get("extracted_text", "")
                extraction_method = file_metadata.metadata.get("extraction_method", "unknown")
                page_count = file_metadata.metadata.get("page_count")
            else:
                # Extract text for preview
                extraction_result = await self.extract_text_content(file_id)
                full_text = extraction_result.extracted_text
                extraction_method = extraction_result.extraction_method
                page_count = extraction_result.page_count
            
            # Generate preview text
            preview_text = full_text[:max_preview_chars] if full_text else ""
            if len(full_text) > max_preview_chars:
                preview_text += "..."
            
            full_text_available = len(full_text) > max_preview_chars
            
            return FilePreview(
                file_id=file_id,
                filename=file_metadata.filename,
                mime_type=file_metadata.mime_type,
                file_size=file_metadata.file_size,
                preview_text=preview_text,
                full_text_available=full_text_available,
                page_count=page_count,
                metadata={
                    "extraction_method": extraction_method,
                    "full_text_length": len(full_text),
                    "preview_length": len(preview_text)
                }
            )
            
        except Exception as e:
            logger.error(f"Failed to generate file preview for {file_id}: {e}")
            raise
    
    async def get_full_extracted_text(self, file_id: int) -> str:
        """
        Get the full extracted text content for a file.
        
        Args:
            file_id: File identifier
            
        Returns:
            Full extracted text content
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If text extraction failed
        """
        try:
            # Get file metadata
            file_metadata = await self.get_file_metadata(file_id)
            
            # Check if text has already been extracted
            if file_metadata.metadata.get("text_extracted"):
                return file_metadata.metadata.get("extracted_text", "")
            else:
                # Extract text
                extraction_result = await self.extract_text_content(file_id)
                if not extraction_result.success:
                    raise ValueError(f"Text extraction failed: {extraction_result.error_message}")
                return extraction_result.extracted_text
                
        except Exception as e:
            logger.error(f"Failed to get full extracted text for {file_id}: {e}")
            raise